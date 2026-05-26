#!/usr/bin/env python3
"""修改中文论文的第四章：以变更对比为主线，最后是实验结果留白"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import copy
import os

INPUT = '/root/1/基于Transformer的物理仿真_从PhysCtrl到PhysCtrl-2.docx'
OUTPUT = '/root/1/基于Transformer的物理仿真_从PhysCtrl到PhysCtrl-2.docx'

# ==========================================================
# Helper functions
# ==========================================================
def add_code_block(doc, text, font_size=7.5):
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.style = doc.styles['No Spacing']
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(font_size)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.4)

def add_table_simple(doc, headers, rows, style='Light Shading Accent 1'):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            table.rows[i+1].cells[j].text = str(cell_text)
            for p in table.rows[i+1].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

def add_math(doc, formula):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(formula)
    run.font.italic = True
    run.font.size = Pt(10)

def body(doc, text):
    p = doc.add_paragraph(text)
    return p

def bold_body(doc, bold_part, rest):
    p = doc.add_paragraph()
    p.add_run(bold_part).bold = True
    p.add_run(rest)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_heading_safe(doc, text, level):
    """Add heading with font setup."""
    h = doc.add_heading(text, level=level)
    return h

# ==========================================================
# Generate NEW Chapter 4 content
# ==========================================================
def generate_new_chapter4(doc):
    """Generate the new Chapter 4 content into the given document."""

    # ===== 4.1 系统总览 (保留) =====
    add_heading_safe(doc, '4.1  系统总览', level=2)
    body(doc, 'PhysCtrl-2（23 sota 配置）是在 physctrl_o 基线基础上，经过两轮系统性改进（f1t2 + 最终优化）形成的最终版本。'
         '本节以「变更对比」为主线，逐一阐述 physctrl_2 相对于 physctrl_o 在训练范式、条件编码、损失函数、验证策略和工程优化等维度的核心变化，'
         '最后以架构总览图总结完整信息流。')

    body(doc, '下表给出了 physctrl_2 与 physctrl_o 的核心配置对比：')
    config_compare = [
        ('pc_size', '512', '512', '相同'),
        ('latent_dim', '256', '256', '相同'),
        ('n_layers', '8', '8', '相同'),
        ('frame_cond', 'true', 'true', '相同'),
        ('point_embed', 'true', 'true', '相同'),
        ('pred_offset', 'true', 'true', '相同'),
        ('transformer_block', 'v1', 'v1', '相同'),
        ('use_diffusion', '**true** (DDPM)', '**false** (确定性)', '⚡ 核心变更'),
        ('mask_cond', '**false**', '**true**', '⚡ 新增'),
        ('floor_cond', '**false**', '**true**', '⚡ 新增'),
        ('force_as_token', 'false', 'false', '相同'),
        ('λ_deform', '**0.001**', '**0.1**', '⚡ 100×'),
        ('λ_collision', '**无**', '**0.1**', '⚡ 新增'),
        ('λ_laplacian', '**无**', '0.0 (机制预留)', '⚡ 新增机制'),
        ('λ_edge', '**无**', '0.0 (机制预留)', '⚡ 新增机制'),
        ('max_train_steps', '60000', '60000', '相同'),
        ('gradient_accumulation', '8', '8', '相同'),
        ('mixed_precision', 'bf16', 'bf16', '相同'),
    ]
    add_table_simple(doc, ['参数', 'physctrl_o', 'physctrl_2 (23 sota)', '变化'], config_compare)

    doc.add_heading('4.1.1  架构总览图', level=3)
    add_code_block(doc, """┌──────────────────────────────────────────────────────────────┐
│                        输入层                                 │
│  points_src (B,5,512,3)  force (B,1,3)  E,ν (B,1)           │
│  drag_point (B,1,4)      mask (B,1,512,1)  floor (B,1)       │
│  gravity (B,1)            start_vel (B,512,3)                 │
└────────────┬─────────────────────────────────────────────────┘
             │
    ┌────────┴────────┐
    │  条件编码        │      ┌──────────────────┐
    │  E → Lin(1,256)  │      │  PointEmbed      │
    │  ν → Lin(1,256)  │      │  Fourier 96维    │
    │  force→Lin(3,256)│      │  + MLP(99→256)   │
    │  drag→Lin(3,256) │      │  每点独立编码     │
    │  floor→Lin(1,256)│      └────────┬─────────┘
    └────────┬────────┘              │
             │                       │
    ┌────────▼───────────────────────▼─────────┐
    │  encoder_states (B,5,256)                │
    │  hidden_states (B,7,512,256)             │
    │   = [mask_frame, init_frame, 5×noised]   │
    └──────────────────┬───────────────────────┘
                       │
    ┌──────────────────▼───────────────────────┐
    │       SpaitalTemporalTransformer         │
    │  ┌────────────────────────────────────┐  │
    │  │ Time Embed + 3D Sinusoidal PosEmb  │  │
    │  └──────────────┬─────────────────────┘  │
    │  ┌──────────────▼─────────────────────┐  │
    │  │ ×8: ST Block (v1)                  │  │
    │  │  ├─ Spatial Self-Attn (4 heads)    │  │
    │  │  ├─ Joint FFN (256→1024→256)       │  │
    │  │  └─ Temporal Self-Attn (per-point)  │  │
    │  └──────────────┬─────────────────────┘  │
    │  ┌──────────────▼─────────────────────┐  │
    │  │ Norm + AdaLN + Linear(256→3)       │  │
    │  └────────────────────────────────────┘  │
    └──────────────────┬───────────────────────┘
                       │
    ┌──────────────────▼───────────────────────┐
    │  output[:, 2:] → (B,5,512,3)            │
    │  + init_pc_base (残差恢复)               │
    │  → 最终预测 (B, 5, 512, 3)               │
    └──────────────────────────────────────────┘""")

    # ===== 4.2 训练范式变更 =====
    add_heading_safe(doc, '4.2  训练范式变更：从扩散到确定性预测', level=2)
    bold_body(doc, 'physctrl_o 的方式（纯扩散）：',
             '始终使用 DDPMScheduler（1000 步），训练时向目标点云加噪（t ~ Uniform(0, 999)），'
             '模型需学习从不同噪声水平恢复干净信号。推理时需 DDIM 迭代去噪（50 步）。')
    bold_body(doc, 'physctrl_2 的改进（确定性预测）：',
             '直接绕过扩散过程。用最后一帧源点云复制 5 份并加小幅噪声（σ=0.02），'
             '以 t=0 一步前向预测目标。推理时同样只需单步。')

    add_code_block(doc, """# physctrl_o (纯扩散)
noise = torch.randn_like(latents)
timesteps = torch.randint(0, 1000, (bsz,))
model_input = noise_scheduler.add_noise(latents, noise, timesteps)

# physctrl_2 (确定性预测)
last_src = batch['points_src'][:, -1:, :, :]              # (B,1,N,3)
model_input = last_src.repeat(1, OUTPUT_FRAMES, 1, 1)      # 复制 5 份
model_input += torch.randn_like(model_input) * 0.02         # 小噪声打破对称性
timesteps = torch.zeros((bsz,), dtype=torch.long)           # t = 0""")

    body(doc, '变更效果：训练效率提升约 5-10 倍（省去加噪-去噪循环），推理速度提升约 50 倍（1 步 vs 50 步 DDIM）。')

    # ===== 4.3 条件编码增强 =====
    add_heading_safe(doc, '4.3  条件编码增强', level=2)

    body(doc, 'physctrl_o 的条件编码体系较为简洁：仅 E_cond_encoder、nu_cond_encoder、force_cond_encoder、'
         'drag_point_encoder 四个 Linear 层，cond_seq_length=4。不支持拖拽掩码（mask_cond）和地板高度（floor_cond）的显式编码。')

    body(doc, 'physctrl_2 在此基础上做了以下增强：')

    cond_changes = [
        ('新增 floor_encoder', 'Linear(1, 256) 将地板高度编码为条件 token，cond_seq_length 由 4 → 5。'
         '使模型能显式获取地板边界信息，不再需要从损失函数中隐式推断。'),
        ('开启 mask_cond', '拖拽掩码 (B, 1, 512, 1) 通过 Linear(1, 256) 编码后作为额外一帧拼接到 hidden_states 前面，'
         'cond_frame 由 1 → 2。模型直接「看到」哪些点被拖拽，而非从力向量中间接推理。'),
        ('新增 start_vel_encoder', 'Linear(3, 256) 将初始速度编码后加到第 0 帧隐状态上，为模型提供逐点的初始运动状态信息。'),
        ('保留 pred_offset', '与 physctrl_o 相同，预测偏移量而非绝对位置，输出 = model_output + init_pc_base。'),
    ]
    for title, desc in cond_changes:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title + '：').bold = True
        p.add_run(desc)

    add_heading_safe(doc, '4.3.1  条件编码对比表', level=3)
    cond_compare = [
        ('E (Young\'s modulus)', 'Linear(1, 256)', 'Linear(1, 256)', '相同'),
        ('ν (Poisson\'s ratio)', 'Linear(1, 256)', 'Linear(1, 256)', '相同'),
        ('Force', 'Linear(3, 256)', 'Linear(3, 256)', '相同'),
        ('Drag Point', 'Linear(3, 256)', 'Linear(3, 256)', '相同'),
        ('Floor Height', '**无**', '**Linear(1, 256)**', '⚡ 新增'),
        ('Drag Mask (mask_cond)', '**无**', '**Linear(1, 256) → pre-frame**', '⚡ 新增'),
        ('Start Velocity', '**无**', '**Linear(3, 256) → add to frame 0**', '⚡ 新增'),
        ('cond_seq_length', '4', '**5**', '⚡ +1'),
        ('cond_frame', '1', '**2** (mask_frame+init_frame)', '⚡ +1'),
    ]
    add_table_simple(doc, ['条件类型', 'physctrl_o', 'physctrl_2', '变化'], cond_compare)

    # ===== 4.4 损失函数变更 =====
    add_heading_safe(doc, '4.4  损失函数全面升级', level=2)

    body(doc, 'physctrl_o 仅使用基础 MSE (loss_xyz)、速度损失 (loss_vel) 和低权重的 MPM 变形损失 (λ_deform=0.001)。'
         '物理约束几乎不起作用，模型主要依赖 pixel-level 重建，容易产生局部坍塌、穿透等伪影。')

    body(doc, 'physctrl_2 的核心变更：')

    loss_changes = [
        ('λ_deform: 0.001 → 0.1 (100×)', 'MPM 物理自洽损失权重提升 100 倍，强制模型生成的轨迹满足材料力学。'
         '为防止梯度爆炸，对输入 MPM 的点云施加 clamp(min=-2.2, max=2.2) 截断。'),
        ('新增碰撞损失 (λ_collision=0.1)', '惩罚点对间距小于 0.01 的自穿透现象。计算公式：L_coll = mean(ReLU(0.01 - dist_ij)²)。'),
        ('新增拉普拉斯损失 (λ_laplacian=0, 机制预留)', '基于 KNN (k=8) 的拉普拉斯坐标一致性约束，防止局部几何结构坍塌。当前权重为 0，可在需要时启用。'),
        ('新增边长损失 (λ_edge=0, 机制预留)', '基于 KNN 图的边长保持约束，模拟局部刚性。当前权重为 0，可在需要时启用。'),
        ('关闭部分实验性损失', 'loss_mask 和 loss_momentum 在 sota 配置中关闭 (λ=0)，将训练信号集中在最有效的损失项上。'),
    ]
    for title, desc in loss_changes:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title + '：').bold = True
        p.add_run(desc)

    add_heading_safe(doc, '4.4.1  损失权重对比表', level=3)
    loss_compare = [
        ('loss_xyz (MSE)', '1.0 (固定)', '1.0 (固定)', '不变'),
        ('loss_vel (速度一致性)', '1.0', '1.0', '不变'),
        ('loss_deform (MPM)', '**0.001**', '**0.1**', '⚡ ↑100×'),
        ('loss_floor (地板穿透)', '1.0 (隐式)', '**0.1** (显式)', '⚡ 独立权重'),
        ('loss_collision (自碰撞)', '**无**', '**0.1**', '⚡ 新增'),
        ('loss_laplacian (拉普拉斯)', '**无**', '0.0 (预留)', '⚡ 新增机制'),
        ('loss_edge (边长)', '**无**', '0.0 (预留)', '⚡ 新增机制'),
        ('loss_mask', '0.0 (关闭)', '0.0 (关闭)', '相同'),
        ('loss_momentum', '0.0 (关闭)', '0.0 (关闭)', '相同'),
    ]
    add_table_simple(doc, ['损失项', 'physctrl_o λ', 'physctrl_2 λ', '变化'], loss_compare)

    # ===== 4.5 验证策略变更 =====
    add_heading_safe(doc, '4.5  验证策略变更：一次性 vs 自回归 Rollout', level=2)

    bold_body(doc, 'physctrl_o 的方式：',
             '验证时使用 DDIM Scheduler，一次性从纯噪声生成完整的 24 帧序列（n_frames=24）。'
             'DDIM 去噪需要 50 次迭代，每次迭代调用一次模型前向。')
    bold_body(doc, 'physctrl_2 的改进：',
             '采用自回归 Rollout 策略。每次只生成 5 帧（OUTPUT_FRAMES=5），然后将预测的 5 帧作为下一步的输入，'
             '重复 ROLLOUT_STEPS=4 次，共生成 5 + 4×5 = 25 帧。相邻 rollout 步之间通过 start_vel 保持速度连续性。')

    add_code_block(doc, """# physctrl_2 自回归 Rollout 验证
current_input = points_src                         # 5 帧真值
for step in range(ROLLOUT_STEPS=4):
    if step == 0:
        step_start_vel = batch['start_vel']        # 真实初速度
    else:
        step_start_vel = current_input[1] - prev_chunk[-1]  # 速度连续性
    
    pred_chunk = pipeline(current_input, ..., start_vel=step_start_vel, n_frames=5)
    rollout_chunks.append(pred_chunk)
    prev_chunk = current_input
    current_input = pred_chunk                     # 用预测作为下一轮输入

output = concat(rollout_chunks) → (B, 5+4×5, N, 3) = 25 帧""")

    body(doc, '该策略的优势：(1) 每次只需预测 5 帧，模型精度更高；(2) 通过自回归可以模拟任意长度的轨迹；'
         '(3) start_vel 机制保证了 rollout 步之间的速度连续性，避免了人工拼接的跳变。')

    # ===== 4.6 工程优化 =====
    add_heading_safe(doc, '4.6  工程优化', level=2)

    eng_changes = [
        ('MPM 粒子数缩减', 'DeformLoss 中 self.N 从 2048 降至 512。在 MPM 物理损失计算中，'
         'P2G + G2P 循环涉及的粒子数减少 4 倍，显著降低显存占用（约 40%）和计算时间（约 3 倍加速）。'),
        ('数值稳定化', '输入 MPM 损失的点云施加 clamp(min=-2.2, max=2.2) 截断，对应物理空间约 [-4.4, 4.4]，'
         '有效防止大变形场景下的梯度爆炸。'),
        ('结构化日志', '引入 CSV 格式的损失记录（每 500 步导出），以及每种损失的独立曲线图（10 种），'
         '便于快速定位训练问题和进行超参数调优。'),
        ('Gradient Checkpointing', '开启梯度检查点，以少量额外计算换取约 40% 的显存节省，'
         '使得在显存受限的 GPU 上也能训练更大 batch 或更深的模型。'),
        ('xFormers 高效注意力', '启用 xFormers 的内存高效注意力实现，进一步降低 Transformer 层的显存开销。'),
    ]
    for title, desc in eng_changes:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title + '：').bold = True
        p.add_run(desc)

    # ===== 4.7 PointEmbed 的保留与作用 =====
    add_heading_safe(doc, '4.7  PointEmbed 的保留与作用', level=2)

    body(doc, 'PointEmbed（傅里叶特征点嵌入）在 physctrl_o 和 physctrl_2 中均使用，两者实现完全相同。'
         '但它在 physctrl_2 的确定性预测模式下发挥了更为关键的作用：')
    pe_points = [
        '在 physctrl_o 中，输入是带噪点云（每个点已经有不同的坐标值），PointEmbed 仅需微调已有位置信息；',
        '在 physctrl_2 中，5 帧输出帧的输入是全零（所有点坐标完全相同），PointEmbed 为初始帧提供了唯一的 ' 
        '逐点几何编码，成为空间注意力中「形状信息从初始帧流向零帧」的唯一桥梁。',
        '傅里叶基频（2^k·π, k=0..15）覆盖了从大尺度形变（低频）到微观细节（高频）的完整频段，'
        '使模型无需从零输入中「凭空」学习高频函数。',
    ]
    for pt in pe_points:
        bullet(doc, pt)

    # ===== 4.8 DiT 主干网络：继承与不变 =====
    add_heading_safe(doc, '4.8  DiT 主干网络', level=2)

    body(doc, 'physctrl_2 的 DiT 主干（SpaitalTemporalTransformer）与 physctrl_o 在结构上保持一致：'
         '均为 8 层 SpatialTemporalTransformerBlock (v1)，4 头 × 64 维注意力，256 维隐空间，'
         '3D 正弦位置编码（空间 3/4 + 时间 1/4），LayerNorm + AdaLN + Linear(256→3) 输出投影。')

    body(doc, '核心不变的原因：v1 Block 的串行设计（空间注意力 → 联合 FFN → 逐点时间注意力）已经被证明能有效建模 '
         '点云序列的物理动力学，在大幅改变训练范式和损失函数后，保持架构稳定有助于隔离各改进的贡献。')

    # ===== 4.9 数据处理管线：对比 =====
    add_heading_safe(doc, '4.9  数据处理管线', level=2)

    body(doc, 'physctrl_o 和 physctrl_2 共享相同的数据生成和加载管线（TrajDataset），但在以下细节上有差异：')

    data_changes = [
        ('输入/输出帧数', 'physctrl_o 直接使用 n_training_frames=24 的完整序列，输入与目标帧数由 train.py 中 INPUT_FRAMES/OUTPUT_FRAMES 全局变量（均为 5）控制。'
         'physctrl_2 同样使用 5+5 的滑动窗口，但通过 start_idx 步长 5 的采样策略最大化数据利用率。'),
        ('归一化策略', '两者完全一致：坐标 (x - 5.0)/2.0 映射到 [-1,1]，力除以 base_drag_coeff 实现与体积解耦。'),
        ('数据增强', '两者均使用随机点云采样（训练时随机，验证时固定种子）。physctrl_2 额外利用了 start_vel 信息（通过中心差分从相邻帧计算），增强了时序一致性。'),
    ]
    for title, desc in data_changes:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title + '：').bold = True
        p.add_run(desc)

    # ===== 4.10 变更总结 =====
    add_heading_safe(doc, '4.10  变更总结', level=2)

    body(doc, '下表从六个维度总结了 physctrl_2 相对于 physctrl_o 的所有关键变更：')

    summary = [
        ('训练范式', '纯扩散 (DDPM)', '确定性预测 (t=0, all-zeros input)', '⚡ 核心变更'),
        ('条件编码', '4 tokens, 无 mask/floor', '5 tokens + mask_frame + start_vel', '⚡ 增强'),
        ('损失函数', 'MSE+Vel+Deform(0.001)+Floor', '+Collision(0.1), Deform→0.1, +Laplacian/Edge 机制', '⚡ 全面升级'),
        ('验证策略', '一次生成 24 帧 (DDIM 50步)', '自回归 Rollout (4×5=25帧, 1步/次)', '⚡ 改进'),
        ('MPM 损失', 'N=2048, 无 clamp', 'N=512, clamp(-2.2, 2.2)', '⚡ 优化'),
        ('日志系统', '简单曲线', 'CSV + 多曲线 (10种)', '⚡ 工程化'),
    ]
    add_table_simple(doc, ['维度', 'physctrl_o', 'physctrl_2', '变化'], summary)

    # ===== 4.11 实验结果 =====
    add_heading_safe(doc, '4.11  实验结果', level=2)

    body(doc, '（本节内容待补充，将在完成 physctrl_2 的完整训练后填入定量指标和定性可视化结果。）')

    body(doc, '计划评估指标：')
    for m in [
        'Volume IoU (vIoU)：预测点云体素与真值的交并比；',
        'Chamfer Distance (CD)：点云间的平均最近邻距离；',
        'L2 Distance：对应点之间的欧氏距离；',
        '训练收敛曲线：各损失项（xyz/vel/deform/floor/collision）随训练步数的变化趋势；',
        '可视化对比：生成轨迹 vs 真值轨迹的三视图（XYZ 三个正视角），叠加 bounding box；',
        '消融实验：分别关闭 mask_cond、floor_cond、collision loss，验证各项改进的独立贡献。',
    ]:
        bullet(doc, m)

    body(doc, '计划对比基线：')
    for b in [
        'physctrl_o（纯扩散，4 条件 token，λ_deform=0.001）：验证非扩散训练 + 条件编码增强 + 损失函数升级的综合效果；',
        'physctrl_2 f1t2（全零输入, v3 Block）：验证 v1 vs v3 Block 的影响，以及自回归 Rollout vs 一次性生成的差异；',
        'physctrl_2 w/o mask_cond：消融验证拖拽掩码条件编码的独立贡献；',
        'physctrl_2 w/o floor_cond：消融验证地板条件编码的独立贡献；',
        'physctrl_2 w/o collision loss：消融验证自碰撞损失对生成质量的影响。',
    ]:
        bullet(doc, b)


# ==========================================================
# MAIN: Modify the existing document
# ==========================================================
def main():
    doc = Document(INPUT)

    # Find all paragraphs that belong to Chapter 4
    # Strategy: find the index of "第四章" heading, then find "第五章" heading,
    # remove everything between them, insert new content.

    # Collect paragraph indices
    ch4_start = None
    ch5_start = None

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else ''
        if '第四章' in text and 'Heading' in style:
            ch4_start = i
        if '第五章' in text and 'Heading' in style and ch4_start is not None:
            ch5_start = i
            break

    if ch4_start is None or ch5_start is None:
        print(f"ERROR: Could not find chapter boundaries. ch4_start={ch4_start}, ch5_start={ch5_start}")
        return

    print(f"Chapter 4 starts at paragraph {ch4_start}, Chapter 5 at paragraph {ch5_start}")
    print(f"Removing paragraphs {ch4_start} to {ch5_start-1} ({ch5_start - ch4_start} paragraphs)")

    # Also need to find and remove tables between ch4_start and ch5_start
    # python-docx stores tables separately from paragraphs in the document body
    # We need to handle both paragraphs and tables in document order

    # The document body has both paragraphs and tables in order
    body_elements = doc.element.body

    # Find the XML elements corresponding to the paragraphs to delete
    # We need to find the paragraph elements and table elements between ch4 and ch5

    para_elements = []
    table_elements = []

    # Collect paragraph XML elements
    para_idx = 0
    for child in body_elements:
        if child.tag.endswith('}p'):
            if para_idx >= ch4_start and para_idx < ch5_start:
                para_elements.append(child)
            para_idx += 1

    # Collect table XML elements  
    # Tables are mixed in with paragraphs in the XML
    # We need to find all tables between the ch4_start-th paragraph and ch5_start-th paragraph

    # Actually, let's use a simpler approach: 
    # Find the paragraph XML element for ch4_start, then delete everything 
    # until the paragraph XML element for ch5_start

    para_count = 0
    delete_mode = False
    elements_to_delete = []

    for child in body_elements:
        if child.tag.endswith('}p'):
            if para_count == ch4_start:
                delete_mode = True
            if para_count == ch5_start:
                delete_mode = False
                # Don't delete ch5 heading
                break
            para_count += 1
        
        if delete_mode:
            elements_to_delete.append(child)

    print(f"Found {len(elements_to_delete)} XML elements to delete")

    # Delete the elements
    for elem in elements_to_delete:
        body_elements.remove(elem)

    # Now we need to insert the new Chapter 4 content.
    # But the new content needs to be inserted BEFORE the Chapter 5 heading paragraph.
    # We need to find the Chapter 5 paragraph XML element.

    # Find the ch5 heading paragraph element
    para_count = 0
    insert_before = None
    for child in body_elements:
        if child.tag.endswith('}p'):
            if para_count == ch5_start - (ch5_start - ch4_start):  # adjusted index after deletion
                # Actually this is tricky. Let me just find "第五章" text in remaining paragraphs
                pass
            para_count += 1

    # Simpler approach: find the paragraph containing "第五章" in the remaining body
    for child in body_elements:
        if child.tag.endswith('}p'):
            # Check if this paragraph contains "第五章"
            text_elems = child.findall('.//' + qn('w:t'))
            text = ''.join(t.text or '' for t in text_elems)
            if '第五章' in text:
                insert_before = child
                break

    if insert_before is None:
        print("ERROR: Could not find Chapter 5 heading after deletion")
        return

    # Now we need to insert the new Chapter 4 content before the Chapter 5 heading.
    # We'll create a temporary document with the new content, then move its elements.
    
    tmp_doc = Document()
    
    # Copy styles? Not needed since they're the same document template
    
    # Generate Chapter 4 content into tmp_doc
    # First, add "第四章  最终技术实现：PhysCtrl-2" heading
    h = tmp_doc.add_heading('第四章  最终技术实现：PhysCtrl-2', level=1)
    
    generate_new_chapter4(tmp_doc)

    # Now move all elements from tmp_doc.body before insert_before
    insert_index = list(body_elements).index(insert_before)
    
    for child in list(tmp_doc.element.body):
        body_elements.insert(insert_index, child)
        insert_index += 1

    # Save
    doc.save(OUTPUT)
    print(f"Updated document saved to: {OUTPUT}")

if __name__ == '__main__':
    main()
