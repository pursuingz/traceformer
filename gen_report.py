#!/usr/bin/env python3
"""生成 PhysCTRL-2 模型结构与 Pipeline 详解的 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_code_block(doc, text, font_size=8):
    """添加代码块"""
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.style = doc.styles['No Spacing']
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(font_size)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)

def create_document():
    doc = Document()
    
    # ===== 样式设置 =====
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # ===== 标题 =====
    title = doc.add_heading('PhysCTRL-2 模型结构与 Pipeline 详解', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('基于 DiT 的物理感知点云形变预测模型')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # spacer
    
    # ===== 目录概览 =====
    doc.add_heading('目录', level=1)
    toc_items = [
        '一、项目概览',
        '二、核心配置参数解读',
        '三、模型结构详解',
        '  3.1 PointEmbed：傅里叶特征点嵌入',
        '  3.2 条件编码器',
        '  3.3 帧条件机制',
        '  3.4 DiT 核心：SpaitalTemporalTransformer',
        '  3.5 SpatialTemporalTransformerBlock 内部结构',
        '四、数据处理 Pipeline',
        '五、训练 Pipeline 详解',
        '六、推理 Pipeline',
        '七、架构总结图',
        '八、文生图 Prompt',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.style = doc.styles['No Spacing']
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    
    doc.add_page_break()
    
    # ===== 一、项目概览 =====
    doc.add_heading('一、项目概览', level=1)
    p = doc.add_paragraph()
    p.add_run('PhysCTRL-2').bold = True
    p.add_run(' 是一个')
    p.add_run('基于 DiT (Diffusion Transformer) 的物理仿真模型').bold = True
    p.add_run('，用于直接在点云上预测物体在受力下的形变与运动轨迹。它基于 CogVideoX 的 Transformer 架构改造而来，引入时空注意力机制来处理 4D 点云数据。')
    
    h = doc.add_heading('核心特征', level=2)
    features = [
        ('模型规模', '约 10M 可训练参数 (base 配置)'),
        ('输入', '点云 (512×3)、力向量 (1×3)、物理参数 (E, ν)、拖拽掩码、地板高度'),
        ('输出', '预测的 5 帧未来点云序列 (512×3×5)'),
        ('训练模式', '无扩散单步前向预测 (non-diffusion direct prediction)'),
        ('验证方式', '自回归 Rollout (4 步，共 25 帧)'),
        ('核心创新', 'Fourier 特征点嵌入 + 时空分离注意力 + 可微分 MPM 物理损失'),
    ]
    table = doc.add_table(rows=len(features)+1, cols=2)
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = '属性'
    hdr[1].text = '描述'
    for i, (k, v) in enumerate(features):
        table.rows[i+1].cells[0].text = k
        table.rows[i+1].cells[1].text = v
    
    doc.add_paragraph()
    
    # ===== 二、核心配置参数解读 =====
    doc.add_heading('二、核心配置参数解读', level=1)
    p = doc.add_paragraph('以下为该 SOTA (23) 配置的关键参数：')
    
    config_params = [
        ('model_type', 'dit_st', '使用 DiT 时空模型'),
        ('pc_size', '512', '点云采样 512 个点'),
        ('model_config.n_layers', '8', 'Transformer 层数'),
        ('model_config.latent_dim', '256', '隐空间维度'),
        ('model_config.frame_cond', 'true', '拼接初始帧作为条件'),
        ('model_config.point_embed', 'true', '使用 Fourier 特征点嵌入'),
        ('model_config.mask_cond', 'true', '将拖拽掩码作为附加帧拼入'),
        ('model_config.pred_offset', 'true', '预测残差偏移量（非绝对位置）'),
        ('model_config.floor_cond', 'true', '编码地板高度'),
        ('model_config.force_as_token', 'false', '力作为 cond token 统一处理'),
        ('model_config.force_as_latent', 'false', '力不拼接到点特征'),
        ('model_config.transformer_block', 'SpatialTemporalTransformerBlock', 'v1 串行 Transformer Block'),
        ('use_diffusion', 'false', '不使用扩散过程，单步前向预测'),
        ('n_frames_interval', '1', '逐帧密集采样'),
        ('lambda_deform', '0.1', '形变物理损失权重'),
        ('lambda_collision', '0.1', '自碰撞损失权重'),
        ('lambda_floor', '0.1', '地板穿透损失权重'),
        ('lambda_vel', '1.0', '速度连续性损失权重'),
        ('gradient_accumulation_steps', '8', '梯度累积 8 步，等效 batch_size=8'),
        ('mixed_precision', 'bf16', 'BFloat16 混合精度训练'),
        ('learning_rate', '1e-4', 'AdamW 学习率'),
        ('max_train_steps', '60000', '最大训练步数'),
    ]
    
    table = doc.add_table(rows=len(config_params)+1, cols=3)
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '参数'
    hdr[1].text = '值'
    hdr[2].text = '含义'
    for i, (k, v, desc) in enumerate(config_params):
        table.rows[i+1].cells[0].text = k
        table.rows[i+1].cells[1].text = v
        table.rows[i+1].cells[2].text = desc
    
    doc.add_paragraph()
    
    # ===== 三、模型结构详解 =====
    doc.add_heading('三、模型结构详解', level=1)
    
    # 3.1 PointEmbed
    doc.add_heading('3.1 PointEmbed：傅里叶特征点嵌入', level=2)
    p = doc.add_paragraph()
    p.add_run('这是模型的关键创新').bold = True
    p.add_run('，将 3D 点坐标通过傅里叶基函数映射到高维空间：')
    
    points = [
        '基频: eₖ = 2ᵏ · π，其中 k = 0, 1, ..., 15（16 个频率）',
        '投影: 对每个点 (x,y,z)，计算 sin(eₖ·x), cos(eₖ·x)，每轴 32 维',
        '三个轴共计 96 维 Fourier 特征',
        '拼接原始坐标 3 维 → 99 维输入',
        'MLP: Linear(99, 256) → 投影到隐空间',
    ]
    for pt in points:
        doc.add_paragraph(pt, style='List Bullet')
    
    p = doc.add_paragraph('该设计使模型能自然编码高频空间细节，类似 NeRF 中著名的 Positional Encoding。')
    
    # 3.2 条件编码器
    doc.add_heading('3.2 条件编码器 (Condition Encoders)', level=2)
    p = doc.add_paragraph('条件序列长度 (cond_seq_length) 动态计算为 5，每个物理条件通过小型线性层映射到 256 维隐空间：')
    
    cond_table = [
        ('E (Young\'s modulus)', 'Linear(1, 256)', '1 token'),
        ('ν (Poisson\'s ratio)', 'Linear(1, 256)', '1 token'),
        ('Force', 'Linear(3, 256)', '1 token'),
        ('Drag Point', 'Linear(3, 256)', '1 token'),
        ('Floor Height', 'Linear(1, 256)', '1 token'),
    ]
    table = doc.add_table(rows=len(cond_table)+1, cols=3)
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '条件类型'
    hdr[1].text = '编码器'
    hdr[2].text = 'Token 数'
    for i, (a, b, c) in enumerate(cond_table):
        table.rows[i+1].cells[0].text = a
        table.rows[i+1].cells[1].text = b
        table.rows[i+1].cells[2].text = c
    
    # 3.3 帧条件
    doc.add_heading('3.3 帧条件机制 (Frame Conditioning)', level=2)
    p = doc.add_paragraph('将初始点云的最后一帧拼接到输入之前，同时将拖拽掩码作为额外条件帧：')
    code = """x_cond = concat([mask_frame, init_pc_last_frame, noisy_input], dim=1)
# 形状: (B, 1+1+5, 512, 3) = (B, 7, 512, 3)
# 经过 PointEmbed → (B, 7, 512, 256)"""
    add_code_block(doc, code)
    
    # 3.4 DiT 核心
    doc.add_heading('3.4 DiT 核心：SpaitalTemporalTransformer', level=2)
    p = doc.add_paragraph('继承自 HuggingFace Diffusers 的 ModelMixin，实现完整 DiT 流程：')
    
    steps = [
        ('时间嵌入', 'timesteps → TimeEmbedding → emb (256-dim)'),
        ('位置编码', '3D Sinusoidal Positional Encoding — 空间 192维 (3/4) + 时间 64维 (1/4)'),
        ('序列拼接', '[encoder(5), hidden(7×512)] → (B, 5+3584, 256)'),
        ('×8 Block', '8 层 SpatialTemporalTransformerBlock (开启 gradient checkpointing)'),
        ('输出投影', 'Final LayerNorm + AdaLayerNorm + Linear(256→3)'),
        ('截取恢复', 'output[:, 2:] → (B, 5, 512, 3)，+ init_pc_base 恢复偏移'),
    ]
    
    table = doc.add_table(rows=len(steps)+1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '步骤'
    hdr[1].text = '操作'
    for i, (s, d) in enumerate(steps):
        table.rows[i+1].cells[0].text = s
        table.rows[i+1].cells[1].text = d
    
    # 3.5 TransformerBlock
    doc.add_heading('3.5 SpatialTemporalTransformerBlock 内部结构 (v1)', level=2)
    
    block_steps = [
        ('Step 1 — 空间自注意力', [
            'CogVideoXLayerNormZero (AdaIN 时间调制)',
            'Multi-Head Attention: 512 个点 × 所有帧 joint attend',
            '注意力头: 4 heads × 64 dim = 256 维',
            'Residual + gate 调制',
        ]),
        ('Step 2 — 联合前馈', [
            '拼接 encoder_hidden_states 和 hidden_states',
            'FeedForward: 256 → 1024 → 256 (GELU)',
            '分离回 hidden 和 encoder，residual + gate',
        ]),
        ('Step 3 — 时间自注意力 (Per-Point)', [
            'Rearrange: (B·F, N, C) → (B·N, F, C)',
            '每个物理点独立在时间轴上 self-attention',
            'AdaLayerNorm 调制 + Residual',
            'Rearrange 回去',
        ]),
    ]
    
    for title, items in block_steps:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('关键设计').bold = True
    p.add_run('：时间注意力在 per-point 维度上独立进行，同一物理点在多帧间的演化被独立建模，不同点之间在时间轴上不互相干扰。这天然适合点云序列，因为同一点的物理演化轨迹是最核心的信息。')
    
    # ===== 四、数据处理 Pipeline =====
    doc.add_heading('四、数据处理 Pipeline', level=1)
    
    doc.add_heading('4.1 H5 文件结构', level=2)
    p = doc.add_paragraph('每个 .h5 文件包含一个完整的 MPM 物理仿真序列：')
    
    h5_fields = [
        ('x', '(T, N, 3)', '点云坐标序列'),
        ('drag_force', '(3,)', '施加的力向量'),
        ('drag_point', '(4,)', '拖拽点位置 + 拖拽点数'),
        ('drag_mask', '(N,)', '哪些点被拖拽'),
        ('E', '()', 'Young\'s modulus'),
        ('nu', '()', 'Poisson\'s ratio'),
        ('vol', '(N,)', '每个粒子的体积'),
        ('F', '(T-1, N, 9)', '形变梯度矩阵 (3×3 展平)'),
        ('C', '(T-1, N, 9)', '仿射速度场'),
        ('gravity', '()', '有无重力'),
        ('floor_height', '()', '地板高度'),
    ]
    table = doc.add_table(rows=len(h5_fields)+1, cols=3)
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '字段'
    hdr[1].text = '形状'
    hdr[2].text = '含义'
    for i, (a, b, c) in enumerate(h5_fields):
        table.rows[i+1].cells[0].text = a
        table.rows[i+1].cells[1].text = b
        table.rows[i+1].cells[2].text = c
    
    doc.add_heading('4.2 滑动窗口采样策略', level=2)
    code = """required_span = (5 + 5 - 1) × 1 + 1 = 10  # 需要 10 帧连续数据

for start_idx in range(0, max_start+1, 5):  # 每隔 5 帧取一个窗口
    input_indices  = [start_idx, start_idx+1, ..., start_idx+4]   # 5 帧输入
    output_indices = [start_idx+5, ..., start_idx+9]              # 5 帧输出"""
    add_code_block(doc, code)
    
    doc.add_heading('4.3 数据归一化', level=2)
    points = [
        '坐标归一化: x_norm = (x - 5.0) / 2.0 — 将 MPM 仿真空间 (±5) 映射到 [-1, 1]',
        '力标准化: force_scaled = force_raw / base_drag_coeff — 使力与粒子体积解耦',
        '点云采样: 如果原始点数 > 512，随机采样 512 个点（训练时随机，验证时固定）',
    ]
    for pt in points:
        doc.add_paragraph(pt, style='List Bullet')
    
    # ===== 五、训练 Pipeline =====
    doc.add_heading('五、训练 Pipeline 详解', level=1)
    
    doc.add_heading('5.1 训练循环流程', level=2)
    code = """每步数据流:
1. batch = dataloader 取一批数据
   latents = batch['points_tgt']  # (B, 5, 512, 3) GT 目标

2. 构建 model_input (无扩散模式):
   last_src_frame = points_src[:, -1]              # (B, 1, 512, 3)
   model_input = last_src_frame.repeat(1, 5, 1, 1)  # 复制 5 份
   model_input += randn * 0.02                      # 小噪声打破对称性

3. 条件 dropout (CFG 训练):
   null_emb = (random > 0).float()  # dropout_rate=0, 全为 1

4. 模型前向:
   pred = model(model_input, timesteps=0, points_src, force, E, nu,
                mask, drag_point, floor_height, gravity, coeff,
                y=None, null_emb=null_emb, start_vel)

5. 计算加权损失 → 反向传播 → 梯度累积(8步) → 更新"""
    add_code_block(doc, code)
    
    doc.add_heading('5.2 损失函数总览', level=2)
    loss_params = [
        ('XYZ Loss (基础)', '1.0 (固定)', 'MSE(pred, target)', '基础位置重建'),
        ('Velocity Loss', '1.0', 'MSE(pred_vel, target_vel)', '速度连续性约束'),
        ('Deformation Loss', '0.1', '可微分 MPM 自洽损失', '物理一致性'),
        ('Floor Loss', '0.1', 'ReLU(floor - pred_y)²', '防穿透地板'),
        ('Collision Loss', '0.1', 'ReLU(0.01 - dist)²', '自碰撞防止'),
        ('Mask Loss', '0.0 (关闭)', 'MSE(pred[mask], target[mask])', '拖拽区域精度'),
        ('Laplacian Loss', '0.0 (关闭)', 'KNN 邻域平滑', '局部光滑'),
        ('Edge Loss', '0.0 (关闭)', 'KNN 边长度保持', '局部刚性'),
        ('Momentum Loss', '0.0 (关闭)', 'MSE(p_int, p_ext)', '动量守恒'),
    ]
    table = doc.add_table(rows=len(loss_params)+1, cols=4)
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '损失项'
    hdr[1].text = '权重 λ'
    hdr[2].text = '公式'
    hdr[3].text = '物理意义'
    for i, (a, b, c, d) in enumerate(loss_params):
        table.rows[i+1].cells[0].text = a
        table.rows[i+1].cells[1].text = b
        table.rows[i+1].cells[2].text = c
        table.rows[i+1].cells[3].text = d
    
    doc.add_heading('5.3 Deformation Loss（可微分 MPM）', level=2)
    p = doc.add_paragraph('这是最复杂的物理损失，将完整的 MPM (Material Point Method) 仿真器嵌入训练循环：')
    steps = [
        '① 反归一化: x = pred × 2 + 5 (恢复 MPM 物理尺度)',
        '② 速度估计: v = (x[t+2] - x[t]) / (2 · dT · 2)',
        '③ P2G (Particle to Grid): 125³ 网格，dx=0.08，Quadratic B-spline 权重 (3×3×3=27 邻域)，grid_m 和 grid_v 插值',
        '④ G2P (Grid to Particle): F_pred = (I + Σ grid_v ⊗ ∇w · dT) @ F_current',
        '⑤ Loss = L1(F_pred, F_gt_next) — 形变梯度预测与真值的 L1 距离',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Bullet')
    
    doc.add_heading('5.4 优化器与训练配置', level=2)
    opt_config = [
        ('优化器', 'AdamW (lr=1e-4, β₁=0.9, β₂=0.999, weight_decay=1e-2)'),
        ('学习率调度', 'Cosine Annealing + 100 步 Warmup'),
        ('混合精度', 'BFloat16 (via Accelerate)'),
        ('梯度累积', '8 步 → 等效 batch_size = 8'),
        ('梯度裁剪', 'max_norm = 1.0'),
        ('最大训练步数', '60,000'),
        ('Checkpoint', '每 2,500 步保存一次，保留最近 10 个'),
        ('验证间隔', '每 500 步'),
        ('Gradient Checkpointing', '开启（节省显存）'),
        ('xFormers 高效注意力', '开启'),
        ('TF32', '开启（Ampere GPU 加速）'),
    ]
    table = doc.add_table(rows=len(opt_config)+1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '配置项'
    hdr[1].text = '值'
    for i, (a, b) in enumerate(opt_config):
        table.rows[i+1].cells[0].text = a
        table.rows[i+1].cells[1].text = b
    
    doc.add_heading('5.5 验证：自回归 Rollout', level=2)
    code = """验证流程（每 500 步执行一次）:
1. 取验证集第一组 5 帧真值输入
2. 模型预测 5 帧输出
3. 以预测的 5 帧作为下一步输入
4. start_vel = 当前输入帧1 - 上一步输出帧-1（速度连续性）
5. 重复 ROLLOUT_STEPS=4 次 → 共 5 + 4×5 = 25 帧
6. 与真值序列对比，生成可视化视频 (XYZ 三视角)"""
    add_code_block(doc, code)
    
    # ===== 六、推理 Pipeline =====
    doc.add_heading('六、推理 Pipeline (pipeline_traj.py)', level=1)
    code = """无扩散模式（use_diffusion=false）:
1. 初始化: sample = init_pc_last_frame.repeat(1, n_frames, 1, 1) + randn×0.02
2. 单步推理: t = 0（timestep 恒为零）
3. model(sample, t=0, init_pc, force, E, nu, mask, drag_point, floor, gravity, coeff, ...)
4. 输出 → 预测的 n_frames 帧点云偏移 + init_pc_base（残差恢复）

支持 Classifier-Free Guidance (guidance_scale > 1.0):
output = output_uncond + guidance_scale × (output_cond - output_uncond)"""
    add_code_block(doc, code)
    
    # ===== 七、架构总结图 =====
    doc.add_heading('七、架构总结图', level=1)
    p = doc.add_paragraph('以下是模型的完整信息流架构总结：')
    
    code = """┌──────────────────────────────────────────────────────────────┐
│                        输入层                                 │
│  points_src (B,5,512,3)  force (B,1,3)  E,ν (B,1)           │
│  drag_point (B,1,4)      mask (B,1,512,1)  floor (B,1)       │
│  gravity (B,1)            start_vel (B,512,3)                 │
└────────────┬─────────────────────────────────────────────────┘
             │
    ┌────────┴────────┐
    │  条件编码        │      ┌──────────────────┐
    │  E → Linear(1,256)│     │  PointEmbed      │
    │  ν → Linear(1,256)│     │  Fourier 96维    │
    │  force→Lin(3,256) │     │  + MLP(99→256)   │
    │  drag→Lin(3,256)  │     │  每点独立编码     │
    │  floor→Lin(1,256) │     └────────┬─────────┘
    └────────┬────────┘              │
             │                       │
    ┌────────▼───────────────────────▼─────────┐
    │  encoder_hidden_states (B, 5, 256)       │
    │  hidden_states (B, 7, 512, 256)          │
    │    = [mask_frame, init_frame, 5×noised]  │
    └──────────────────┬───────────────────────┘
                       │
    ┌──────────────────▼───────────────────────┐
    │       SpaitalTemporalTransformer         │
    │  ┌────────────────────────────────────┐  │
    │  │ Time Embed + 3D Sinusoidal PosEmb  │  │
    │  └──────────────┬─────────────────────┘  │
    │                 │                         │
    │  ┌──────────────▼─────────────────────┐  │
    │  │ ×8: SpatialTemporalTransformerBlock│  │
    │  │  ├─ Spatial Self-Attn (4 heads)    │  │
    │  │  ├─ Joint FFN (256→1024→256)       │  │
    │  │  └─ Temporal Self-Attn (per-point)  │  │
    │  └──────────────┬─────────────────────┘  │
    │                 │                         │
    │  ┌──────────────▼─────────────────────┐  │
    │  │ Norm + AdaLN + Linear(256→3)       │  │
    │  └────────────────────────────────────┘  │
    └──────────────────┬───────────────────────┘
                       │
    ┌──────────────────▼───────────────────────┐
    │  截取 output[:, 2:] → (B, 5, 512, 3)    │
    │  + init_pc_base (残差恢复)               │
    │  → 最终预测 (B, 5, 512, 3)               │
    └──────────────────────────────────────────┘"""
    add_code_block(doc, code)
    
    # ===== 八、文生图 Prompt =====
    doc.add_heading('八、文生图 Prompt', level=1)
    p = doc.add_paragraph('以下 prompt 可用于将上述架构总结图转换为专业的技术示意图：')
    
    doc.add_heading('Prompt (English)', level=2)
    prompt_en = """A professional deep learning architecture diagram in clean technical illustration style, white background. The diagram shows a neural network pipeline for physics-based point cloud deformation prediction.

Top section labeled "INPUTS": multiple colored boxes showing (1) Point Cloud Sequence [B,5,512,3] in blue, (2) Force Vector [B,1,3] in red, (3) Material Properties E,nu in green, (4) Drag Mask & Point in orange, (5) Floor Height in purple, (6) Gravity & Start Velocity in teal.

Two parallel encoding branches below:
Left branch "CONDITION ENCODERS": 5 small linear layers (E→Linear, nu→Linear, Force→Linear, Drag→Linear, Floor→Linear) merging into encoder_hidden_states [B,5,256].
Right branch "PointEmbed": showing Fourier feature encoding with sine/cosine wave symbols, projecting 3D points through frequencies 2^k*pi, then MLP(99→256).

Both branches feed into a large central block "SpaitalTemporalTransformer (DiT Core)":
- Top sub-block: "Time Embedding + 3D Sinusoidal Positional Encoding" 
- Middle sub-block: "×8 Stacked Blocks" each containing three components side by side: "Spatial Self-Attention (4 heads)" → "Joint FFN (256→1024→256)" → "Temporal Self-Attention (per-point)"
- Bottom sub-block: "LayerNorm + AdaLN Modulation + Linear(256→3)"

Final output section: "Slice condition frames → Residual add init_pc → Output [B,5,512,3]".

Use clean rectangular boxes with rounded corners, directional arrows showing data flow, consistent color coding (blue for point data, red for forces, green for material, orange for drag, purple for floor). Professional academic paper style, minimal text labels, clear hierarchy. Vector diagram aesthetic, suitable for a computer vision / machine learning conference paper."""
    add_code_block(doc, prompt_en)
    
    doc.add_heading('Prompt (中文补充说明)', level=2)
    prompt_cn = """一张专业的深度学习架构示意图，白色背景，学术论文风格。展示一个用于物理感知点云形变预测的神经网络 Pipeline。

顶部"输入层"：多个彩色方框分别标注 点云序列(B,5,512,3)、力向量(B,1,3)、材料参数E,ν、拖拽掩码与点、地板高度、重力与初速度。

分为两个并行的编码分支：
左侧"条件编码器"分支：5 个 Linear 层将物理参数映射到 [B,5,256]。
右侧"PointEmbed"分支：展示傅里叶特征嵌入，用正弦/余弦波形符号表示频率 2^k·π 的投影过程，然后 MLP(99→256)。

两个分支汇入中央大框"SpaitalTemporalTransformer (DiT 核心)"：
- 顶部子框：时间嵌入 + 3D 正弦位置编码
- 中间子框：×8 堆叠的 Block，每个 Block 包含三个组件：空间自注意力(4头) → 联合FFN(256→1024→256) → 时间自注意力(逐点)
- 底部子框：LayerNorm + AdaLN调制 + Linear(256→3)

最终输出：截取条件帧 → 残差加回初始坐标 → 输出(B,5,512,3)

使用圆角矩形框，方向箭头表示数据流，统一配色（蓝色=点数据，红色=力，绿色=材料，橙色=拖拽，紫色=地板），专业学术风格，最少文字标注，清晰的层级关系。"""
    add_code_block(doc, prompt_cn)
    
    # ===== 保存 =====
    output_path = '/root/1/physctrl_2/PhysCTRL2_模型结构与Pipeline详解.docx'
    doc.save(output_path)
    print(f'文档已保存到: {output_path}')
    return output_path

if __name__ == '__main__':
    create_document()
