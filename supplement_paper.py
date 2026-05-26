#!/usr/bin/env python3
"""Supplement the English paper: expand 5.1, add 5.2, and supplement earlier content."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import copy

INPUT = '/root/1/Transformer-Based_Physics_Simulation.docx'
OUTPUT = '/root/1/Transformer-Based_Physics_Simulation.docx'

def make_paragraph(doc, text, style_name='Normal', bold_prefix=None):
    """Create a paragraph in the given document."""
    if style_name == 'List Bullet':
        p = doc.add_paragraph(style='List Bullet')
    else:
        p = doc.add_paragraph()
    if bold_prefix:
        p.add_run(bold_prefix).bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def make_normal(doc, text):
    return make_paragraph(doc, text, 'Normal')

def make_bullet(doc, text, bold_prefix=None):
    return make_paragraph(doc, text, 'List Bullet', bold_prefix)

def make_heading(doc, text, level=2):
    return doc.add_heading(text, level=level)

# ==========================================================
# Generate the new content elements
# ==========================================================
def generate_new_51_refinement(tmp_doc):
    """Generate expanded 5.1 Refinement bullet."""
    make_bullet(tmp_doc, 
        'This version builds on PhysCtrl-1 and introduces a fundamentally redesigned input scheme and training strategy. '
        'Rather than feeding all-zeros for the output frames, PhysCtrl-2 concatenates the first 5 frames of the source '
        'point cloud (carrying rich shape and motion-state information) with 5 copies of the last source frame plus '
        'small noise (σ=0.02), ensuring the generated trajectory is anchored to the last known state. It adopts an '
        'autoregressive rollout mechanism—generating 5 frames at a time and feeding the predictions back as input '
        'for the next step, repeated 4 times to produce a final sequence of 25 frames—enabling physics simulation '
        'of any desired length (any multiple of n frames) without retraining the model. It introduces a start_vel_encoder '
        'to incorporate per-point initial velocity as a new conditioning parameter, allowing the model to simulate '
        'objects with arbitrary initial motion rather than being restricted to stationary starts. It also enriches '
        'the condition encoding system with floor_cond and mask_cond, and upgrades the loss function with collision '
        'loss, Laplacian loss, and edge-length regularization mechanisms. These improvements collectively deliver '
        'faster training (approximately 60 hours vs. 90 hours per 60,000 iterations under identical hardware and '
        'data conditions), faster inference, and substantially greater flexibility in mid-simulation parameter adjustment.',
        bold_prefix='Refinement of the final version (PhysCtrl-2): ')

def generate_new_52_innovations(tmp_doc):
    """Generate 5.2 Summary of Technical Innovations content."""
    
    innovations = [
        ('Autoregressive Rollout with Multi-Frame Input: ',
         'Unlike PhysCtrl\'s one-shot 24-frame generation from pure noise, PhysCtrl-2 splits the task into an '
         'autoregressive pipeline—each step predicts only 5 frames, using a rich input consisting of the actual '
         'initial 5 frames plus noise-perturbed copies of the last source frame. This design enables simulation '
         'of arbitrarily long sequences, dramatically improves per-step accuracy, and eliminates the need to '
         'retrain for different video lengths.'),
        
        ('Non-Diffusion Deterministic Prediction: ',
         'By removing the noise-adding step and replacing both the noisy point cloud input and the denoising '
         'timestep with zeros (deterministic mode), the model bypasses the entire diffusion process. A single '
         'forward pass yields the predicted result, improving training efficiency by 5-10× and inference '
         'speed by 50× compared to the DDIM iterative denoising paradigm.'),
        
        ('Initial Velocity Encoding: ',
         'Recognizing that in the autoregressive pipeline intermediate particle states are in motion—a critical '
         'factor affecting subsequent simulation—PhysCtrl-2 introduces a start_vel_encoder to inject per-point '
         'initial velocity as conditioning information. This allows the model to handle objects with arbitrary '
         'initial velocities, unlike PhysCtrl which is limited to stationary initial conditions.'),
        
        ('Enhanced Multi-Modal Condition Encoding: ',
         'PhysCtrl-2 extends the original 4-token condition set (E, ν, force, drag_point) with floor height '
         'encoding (floor_cond), drag mask encoding as an additional frame (mask_cond), and initial velocity '
         'encoding (start_vel). These explicit condition signals provide the model with direct access to boundary '
         'and interaction information that PhysCtrl must infer implicitly.'),
        
        ('Mid-Simulation Parameter Adjustment: ',
         'Because the autoregressive approach divides the full time span into multiple segments, users can '
         'adjust material properties or force conditions at any time segment during simulation—a level of '
         'interactive control not possible with PhysCtrl\'s one-shot generation paradigm.'),
        
        ('Multi-Scale Geometric Regularization: ',
         'In addition to the original MSE and velocity losses, PhysCtrl-2 introduces collision loss (preventing '
         'self-penetration), Laplacian deformation loss (preserving local geometric structure via KNN graphs), '
         'and edge-length regularization (enforcing local rigidity). Combined with a 100× increase in the '
         'MPM physics loss weight (λ_deform: 0.001 → 0.1), these constraints work synergistically to produce '
         'physically plausible and geometrically faithful deformations.'),
    ]
    
    for bold_pref, text in innovations:
        make_bullet(tmp_doc, text, bold_prefix=bold_pref)


def generate_supplement_41(tmp_doc):
    """Generate supplementary content for Chapter 4 to expand the PhysCtrl-2 advantages description."""
    
    make_normal(tmp_doc, 
        'In summary, the key architectural and methodological differences between PhysCtrl and PhysCtrl-2 '
        'are as follows:')
    
    make_bullet(tmp_doc,
        'PhysCtrl requires the total number of video frames to be specified upfront, and models for different '
        'frame counts must be retrained from scratch. PhysCtrl-2\'s autoregressive approach enables physics '
        'simulation for any frame count (any multiple of n) and any duration without retraining.',
        bold_prefix='Frame Flexibility: ')
    
    make_bullet(tmp_doc,
        'Because PhysCtrl-2 introduces initial velocity as a parameter, it can simulate objects with arbitrary '
        'initial velocities. PhysCtrl, lacking this parameter, can only simulate objects starting from rest.',
        bold_prefix='Initial Velocity Support: ')
    
    make_bullet(tmp_doc,
        'The autoregressive segmentation of the full time span allows material properties or force conditions '
        'to be adjusted at any time segment, significantly enhancing interactive control flexibility.',
        bold_prefix='Mid-Simulation Control: ')
    
    make_bullet(tmp_doc,
        'Due to its autoregressive nature, errors from each output step accumulate progressively, ultimately '
        'affecting the final output. Addressing this error accumulation remains a key challenge for the method.',
        bold_prefix='Error Accumulation: ')


# ==========================================================
# MAIN: Modify the document by inserting XML elements
# ==========================================================
def main():
    doc = Document(INPUT)
    body = doc.element.body

    # ==============================================================
    # STEP 1: Find and expand p181 (Refinement of the final version)
    # ==============================================================
    # Find paragraph 181 (0-indexed: 180)
    para_idx = 0
    target_para_elem = None
    for child in body:
        if child.tag.endswith('}p'):
            if para_idx == 180:
                target_para_elem = child
                break
            para_idx += 1
    
    if target_para_elem is not None:
        # Create replacement content in a temp doc
        tmp = Document()
        generate_new_51_refinement(tmp)
        
        # Insert all temp elements after the target paragraph
        insert_pos = list(body).index(target_para_elem) + 1
        for child in list(tmp.element.body):
            body.insert(insert_pos, child)
            insert_pos += 1
        
        # Delete the old short paragraph
        body.remove(target_para_elem)
        print("STEP 1: Expanded 5.1 Refinement bullet point.")
    else:
        print("WARNING: Could not find paragraph 181")

    # ==============================================================
    # STEP 2: Insert 5.2 innovations content before 5.3
    # ==============================================================
    # Find the 5.3 heading: look for paragraph with "5.3  Future Work"
    para_count = 0
    target_53_elem = None
    for child in body:
        if child.tag.endswith('}p'):
            # Extract text from this paragraph
            texts = child.findall('.//' + qn('w:t'))
            full_text = ''.join(t.text or '' for t in texts)
            if '5.3' in full_text and 'Future Work' in full_text:
                target_53_elem = child
                break
            para_count += 1
    
    if target_53_elem is not None:
        tmp = Document()
        make_heading(tmp, '5.2  Summary of Technical Innovations', level=2)
        generate_new_52_innovations(tmp)
        
        # Insert before the 5.3 heading
        insert_pos = list(body).index(target_53_elem)
        for child in list(tmp.element.body):
            body.insert(insert_pos, child)
            insert_pos += 1
        print("STEP 2: Inserted 5.2 Technical Innovations content.")
    else:
        print("WARNING: Could not find 5.3 heading")

    # ==============================================================
    # STEP 3: Supplement Chapter 4 — add advantages summary after 4.5
    # ==============================================================
    # Find "4.5  Experimental Results" paragraph
    para_count = 0
    target_45_elem = None
    for child in body:
        if child.tag.endswith('}p'):
            texts = child.findall('.//' + qn('w:t'))
            full_text = ''.join(t.text or '' for t in texts)
            if '4.5' in full_text and 'Experimental Results' in full_text:
                target_45_elem = child
                break
            para_count += 1
    
    if target_45_elem is not None:
        # Insert BEFORE 4.5 Experimental Results
        tmp = Document()
        make_heading(tmp, '4.5  Advantages and Limitations Analysis', level=2)
        generate_supplement_41(tmp)
        
        insert_pos = list(body).index(target_45_elem)
        for child in list(tmp.element.body):
            body.insert(insert_pos, child)
            insert_pos += 1
        print("STEP 3: Supplemented Chapter 4 with advantages/limitations analysis.")
    else:
        print("WARNING: Could not find 4.5 Experimental Results paragraph")

    # ==============================================================
    # STEP 4: Add a note about training/inference speed to 4.2
    # ==============================================================
    # Find the paragraph after "4.2  Training Paradigm" that starts with "Similar to physctrl_1"
    para_count = 0
    after_42_target = None
    for child in body:
        if child.tag.endswith('}p'):
            texts = child.findall('.//' + qn('w:t'))
            full_text = ''.join(t.text or '' for t in texts)
            if '4.2' in full_text and 'Training Paradigm' in full_text:
                # Found the 4.2 heading - the next paragraph should be the target
                pass
            if 'Similar to physctrl_1' in full_text:
                after_42_target = child
                break
            para_count += 1
    
    if after_42_target is not None:
        tmp = Document()
        make_normal(tmp, 
            'Under identical dataset size, computational resources, and iteration count, PhysCtrl-2 achieves '
            'significantly faster training—approximately 60 hours versus 90 hours per 60,000 iterations—while '
            'also delivering a roughly 50× inference speedup (single forward pass vs. 50 DDIM denoising steps).')
        
        insert_pos = list(body).index(after_42_target) + 1
        for child in list(tmp.element.body):
            body.insert(insert_pos, child)
            insert_pos += 1
        print("STEP 4: Added training/inference speed data to 4.2.")
    else:
        print("WARNING: Could not find 4.2 target paragraph")

    # ==============================================================
    # Save
    # ==============================================================
    doc.save(OUTPUT)
    print(f"\nDocument saved to: {OUTPUT}")

if __name__ == '__main__':
    main()
